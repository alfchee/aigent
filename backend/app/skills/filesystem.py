import base64
import io
import json
import mimetypes
from pathlib import Path

from app.core.filesystem import SessionWorkspace, cleanup_artifacts, delete_artifact, restore_artifact
from app.core.runtime_context import emit_event


def get_filesystem_tools(session_id: str):
    workspace = SessionWorkspace(session_id)

    def create_file(filepath: str, content: str, encoding: str = "utf-8") -> str:
        """
        Creates a new file in the session workspace.
        
        Args:
            filepath (str): The path to the file to create (e.g., "notes/todo.txt").
            content (str): The content to write. If encoding is "base64", this must be a base64 encoded string.
            encoding (str): Encoding to use ("utf-8" or "base64"). Defaults to "utf-8".
            
        Returns:
            str: JSON string containing the file metadata if successful, or error message.
        """
        try:
            if encoding.lower() == "base64":
                meta = workspace.write_base64(filepath, content)
            else:
                meta = workspace.write_text(filepath, content, encoding=encoding)
            
            emit_event("artifact", {"session_id": session_id, "op": "write", "path": meta.get("path"), "meta": meta})
            return json.dumps({"saved": meta}, ensure_ascii=False)
        except Exception as e:
            return f"Error creating file: {str(e)}"

    def update_file(filepath: str, start_line: int, end_line: int, new_content: str) -> str:
        """
        Updates a specific range of lines in an existing file.
        
        Args:
            filepath (str): The path to the file to update.
            start_line (int): The 1-based start line number.
            end_line (int): The 1-based end line number. Use 0 to just insert at start_line.
            new_content (str): The new content to replace the lines with.
            
        Returns:
            str: JSON string containing the file metadata, or error message.
        """
        try:
            # Read existing content
            full_text = workspace.read_text(filepath)
            lines = full_text.splitlines(keepends=True)
            
            # Validate indices
            if start_line < 1:
                return "Error: start_line must be >= 1"
            
            # Adjust to 0-based
            start_idx = start_line - 1
            
            # If file is empty or lines are fewer than start_idx, append?
            # For simplicity, let's strictly follow the lines existing.
            if start_idx > len(lines):
                 # Append mode if start_line is beyond end?
                 # Let's just append if it's the next line, otherwise error?
                 if start_idx == len(lines):
                     pass
                 else:
                     return f"Error: start_line {start_line} is out of bounds (file has {len(lines)} lines)"

            # Handle end_line
            if end_line < start_line:
                end_idx = start_idx
            else:
                end_idx = end_line 
            
            # Apply change
            lines[start_idx:end_idx] = [new_content] 
            
            new_text = "".join(lines)
            
            meta = workspace.write_text(filepath, new_text)
            emit_event("artifact", {"session_id": session_id, "op": "update", "path": meta.get("path"), "meta": meta})
            return json.dumps({"saved": meta}, ensure_ascii=False)
            
        except Exception as e:
            return f"Error updating file: {str(e)}"

    def read_file(filepath: str, max_bytes: int = 1_000_000) -> str:
        """
        Reads the content of a file from the session workspace.
        Supports text, binary (as base64), PDF, and DOCX text/table extraction.
        
        Args:
            filepath (str): The path to the file to read.
            max_bytes (int): Maximum bytes to read before truncating. Defaults to 1MB.
            
        Returns:
            str: The content of the file (text or JSON with base64), or error message.
        """
        try:
            filename = Path(str(filepath)).name
            mime_type, _ = mimetypes.guess_type(filename)

            if str(filepath).lower().endswith(".docx"):
                try:
                    import docx
                except ImportError:
                    return "Error reading DOCX: missing dependency python-docx"
                
                try:
                    raw = workspace.read_bytes(filepath)
                    doc = docx.Document(io.BytesIO(raw))
                    
                    content_parts = []
                    content_parts.append("--- Document Content (Text) ---")
                    content_parts.append("\n".join([p.text for p in doc.paragraphs if p.text.strip()]))
                    
                    if doc.tables:
                        content_parts.append("\n--- Document Tables ---")
                        for i, table in enumerate(doc.tables):
                            content_parts.append(f"\n[Table {i+1}]")
                            rows = []
                            for row in table.rows:
                                cells = [str(cell.text).strip().replace('\n', ' ') for cell in row.cells]
                                rows.append("| " + " | ".join(cells) + " |")
                            if rows:
                                content_parts.append(rows[0])
                                content_parts.append("| " + " | ".join(['---'] * len(table.rows[0].cells)) + " |")
                                content_parts.extend(rows[1:])
                    
                    full_text = "\n".join(content_parts)
                    if len(full_text.encode('utf-8')) > max_bytes:
                        return full_text[:max_bytes] + "\n...[Content Truncated]"
                    return full_text
                except Exception as e:
                    return f"Error parsing DOCX: {str(e)}"

            if (mime_type == "application/pdf") or str(filepath).lower().endswith(".pdf"):
                try:
                    from pypdf import PdfReader
                except Exception as e:
                    return f"Error reading PDF: missing dependency pypdf ({str(e)})"
                raw = workspace.read_bytes(filepath)
                if len(raw) > max_bytes:
                    return json.dumps(
                        {"path": filepath, "mime_type": mime_type or "application/pdf", "size_bytes": len(raw), "truncated": True},
                        ensure_ascii=False,
                    )
                reader = PdfReader(io.BytesIO(raw))
                parts: list[str] = []
                for page in reader.pages:
                    try:
                        parts.append(page.extract_text() or "")
                    except Exception:
                        parts.append("")
                return "\n".join(parts).strip()

            raw = workspace.read_bytes(filepath)
            if len(raw) > max_bytes:
                return json.dumps(
                    {"path": filepath, "mime_type": mime_type, "size_bytes": len(raw), "truncated": True},
                    ensure_ascii=False,
                )

            is_text_like = bool(mime_type) and (
                mime_type.startswith("text/") or mime_type in ("application/json", "application/xml")
            )
            is_probably_binary = b"\x00" in raw
            if (mime_type and not is_text_like) or is_probably_binary:
                return json.dumps(
                    {
                        "path": filepath,
                        "mime_type": mime_type,
                        "size_bytes": len(raw),
                        "base64": base64.b64encode(raw).decode("utf-8"),
                    },
                    ensure_ascii=False,
                )

            return raw.decode("utf-8", errors="replace")
        except Exception as e:
            return f"Error reading file: {str(e)}"

    def list_files(directory: str = "/") -> str:
        """
        Lists all files in the session workspace directory.
        
        Args:
            directory (str): The directory to list files from (default is root "/").
            
        Returns:
            str: JSON string containing the list of files, or error message.
        """
        try:
            files = workspace.list_files(directory)
            if not files:
                return "El workspace está vacío."
            return json.dumps({"directory": directory, "files": files}, ensure_ascii=False)
        except Exception as e:
            return f"Error listing files: {str(e)}"

    def delete_file(filepath: str, reason: str = "", actor: str = "agent") -> str:
        try:
            result = delete_artifact(session_id, filepath, actor=actor, reason=reason, allow_archived=False)
            emit_event(
                "artifact",
                {
                    "session_id": session_id,
                    "op": "delete",
                    "path": result.get("path"),
                    "trash_id": result.get("trash_id"),
                    "freed_bytes": result.get("freed_bytes"),
                    "restore_until": result.get("restore_until"),
                },
            )
            return json.dumps(result, ensure_ascii=False)
        except Exception as e:
            return f"Error deleting file: {str(e)}"

    def restore_file(trash_id: str, actor: str = "agent") -> str:
        try:
            result = restore_artifact(session_id, trash_id, actor=actor, allow_archived=False)
            emit_event(
                "artifact",
                {
                    "session_id": session_id,
                    "op": "restore",
                    "path": result.get("path"),
                    "trash_id": trash_id,
                    "size_bytes": result.get("size_bytes"),
                },
            )
            return json.dumps(result, ensure_ascii=False)
        except Exception as e:
            return f"Error restoring file: {str(e)}"

    def cleanup_files(criteria_json: str, actor: str = "agent", dry_run: bool = False) -> str:
        try:
            criteria = json.loads(criteria_json) if criteria_json else {}
            result = cleanup_artifacts(session_id, criteria, actor=actor, allow_archived=False, dry_run=dry_run)
            emit_event(
                "artifact",
                {
                    "session_id": session_id,
                    "op": "cleanup",
                    "count": result.get("count"),
                    "freed_bytes": result.get("freed_bytes"),
                    "dry_run": dry_run,
                },
            )
            return json.dumps(result, ensure_ascii=False)
        except Exception as e:
            return f"Error cleaning files: {str(e)}"

    return [create_file, read_file, update_file, list_files, delete_file, restore_file, cleanup_files]
